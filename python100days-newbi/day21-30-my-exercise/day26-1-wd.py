""" word and ppt"""
from docx import Document
from docx.shared import Cm,Pt
from docx.document import Document as Doc

# create a docment obj
document = Document()
# big title
document.add_heading("Happy Learning Python Word",0)
# paragraph
p = document.add_paragraph('Python是一门非常流行的编程语言，它')
# run obj?
run = p.add_run("Simple")
# set the it to bold
run.bold = True
# font size
run.font.size = Pt(18)
p.add_run('And')
run = p.add_run("Elegant")
run.font.size = Pt(18)
# add an underline
run.underline = True
p.add_run(".")

# level 1
document.add_heading("Heading,Level1",1)
# paragraph with style
document.add_paragraph('Intense quote', style='Intense Quote')
# ul
document.add_paragraph( 'first item in unordered list', style='List Bullet')
document.add_paragraph(
    'second item in unordered list', style='List Bullet'
)
# 添加有序列表
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
document.add_paragraph(
    'second item in ordered list', style='List Number'
)
# add pic
document.add_picture("resources/guido.jpg",width=Cm(5.2))
# section
document.add_section()

records = (
    ("Jack","Male",35,"Programmer"),
    ("Mary","Female",30,"Debugger"),
    ("Frank","Male",25,"Programmer"),
    ("Jade","Female",25,"Software Architect"),
)

# add table
table = document.add_table(rows=1,cols=4)
table.style = 'Light List'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'NAME'
hdr_cells[1].text = 'GENDER'
hdr_cells[2].text = 'AGE'
hdr_cells[2].text = 'JOB'
for name,gender,age,job in records:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = gender
    row_cells[2].text = str(age)
    row_cells[3].text = job
# add page break
document.add_page_break()
# save it
document.save("demo.docx")    
